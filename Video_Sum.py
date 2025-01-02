import os
import shutil
import datetime
import time
import re
import cv2
import openai
from docx import Document
from moviepy.editor import VideoFileClip
from transformers import CLIPProcessor, CLIPModel
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
import pytesseract
import threading
from whisper import import os
import shutil
import datetime
import time
import re
import cv2
import openai
from docx import Document
from moviepy.editor import VideoFileClip
from transformers import CLIPProcessor, CLIPModel
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
import pytesseract
import threading
from whisper import load_model
from pydub import AudioSegment
import numpy as np
import webrtcvad
import wave
import contextlib

# Set your OpenAI API key
openai.api_key = "YOUR_API_KEY_HERE"

# Initialize CLIP model for visual analysis
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
whisper_model = load_model("base")  # Corrected: Load Whisper model for transcription

# Paths
input_path = r"C:\Users\it\Videos\Video Sum\Input"
archive_path = r"C:\Users\it\Videos\Video Sum\Input\Archive"
output_path = r"C:\Users\it\Videos\Video Sum\Output"

# Global variables for the timer
timer_running = False
start_time = 0
total_api_time = 0  # Added to accumulate the total API call time

def sanitize_title(title):
    sanitized_title = re.sub(r'[<>:"/\\|?*]', '', title).strip()
    return sanitized_title

def create_output_folder(title):
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    folder_name = f"{title} - {current_date} - Project"
    folder_path = os.path.join(output_path, folder_name)

    video_chunks_folder = os.path.join(folder_path, "Video Chunks")
    audio_chunks_folder = os.path.join(folder_path, "Audio Chunks")
    os.makedirs(video_chunks_folder, exist_ok=True)
    os.makedirs(audio_chunks_folder, exist_ok=True)

    return folder_path, video_chunks_folder, audio_chunks_folder

def split_video_chunk(chunk, video_output_folder, audio_output_folder, index):
    base_filename = os.path.splitext(os.path.basename(chunk.filename))[0]
    video_chunk_filename = os.path.join(video_output_folder, f"{base_filename}_chunk_{index + 1}.mp4")
    audio_chunk_filename = os.path.join(audio_output_folder, f"{base_filename}_chunk_{index + 1}.mp3")

    chunk = chunk.without_audio()
    chunk.write_videofile(video_chunk_filename, codec="libx264", bitrate="500k")
    chunk.audio.write_audiofile(audio_chunk_filename, codec="mp3", bitrate="128k")

    return video_chunk_filename, audio_chunk_filename

def split_video(input_file, video_output_folder, audio_output_folder, max_duration=600):
    clip = VideoFileClip(input_file)
    duration = clip.duration

    with ThreadPoolExecutor() as executor:
        futures = []
        for i in range(0, int(duration), max_duration):
            chunk = clip.subclip(i, min(i + max_duration, duration))
            futures.append(executor.submit(split_video_chunk, chunk, video_output_folder, audio_output_folder, i // max_duration))

        video_chunks = [future.result() for future in futures]

    return video_chunks

def extract_key_frames(video_path, output_folder, interval=60):
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS)
    frame_interval = int(fps * interval)

    os.makedirs(output_folder, exist_ok=True)
    frame_count = 0
    success, frame = cap.read()
    key_frame_paths = []

    while success:
        if frame_count % frame_interval == 0:
            frame_time = frame_count / fps
            frame_filename = f"frame_{int(frame_time)}.jpg"
            frame_path = os.path.join(output_folder, frame_filename)
            cv2.imwrite(frame_path, frame)
            key_frame_paths.append(frame_path)
        success, frame = cap.read()
        frame_count += 1

    cap.release()
    return key_frame_paths

def extract_text_from_frames(frame_paths):
    extracted_texts = []
    for frame_path in frame_paths:
        image = Image.open(frame_path)
        text = pytesseract.image_to_string(image)
        extracted_texts.append(text)
    return extracted_texts

def analyze_transcript(transcript, extracted_texts):
    procedures = []
    for line in transcript.splitlines():
        if "log in" in line.lower() or "login" in line.lower():
            procedures.append("Step: Log in to VSR system")
        elif "submit" in line.lower():
            procedures.append("Step: Submit the form or data")
        elif "process" in line.lower():
            procedures.append("Step: Process the order or request")
    for text in extracted_texts:
        if "order" in text.lower():
            procedures.append("Step: Review and finalize the order in the VSR system")
    if not procedures:
        procedures.append("No explicit procedures were found in the video.")
    return procedures

def live_timer():
    global timer_running, start_time, total_api_time
    start_time = time.time()
    while timer_running:
        elapsed_time = time.time() - start_time
        print(f"\rTime elapsed: {int(elapsed_time)} seconds", end="")
        time.sleep(1)

def split_text_into_chunks(text, max_tokens=3000):
    sentences = text.split('. ')
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_tokens:
            current_chunk += sentence + ". "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

def filter_audio_segments(audio_chunk_path):
    vad = webrtcvad.Vad()
    vad.set_mode(1)  # Set VAD aggressiveness (0-3). Higher means more aggressive filtering of noise.

    audio = AudioSegment.from_file(audio_chunk_path)
    raw_audio = np.array(audio.get_array_of_samples())

    sample_rate = audio.frame_rate
    num_channels = audio.channels

    window_duration = 0.03  # 30 ms windows
    samples_per_window = int(sample_rate * window_duration * num_channels)

    segments = []

    for i in range(0, len(raw_audio), samples_per_window):
        segment = raw_audio[i:i + samples_per_window]
        if len(segment) < samples_per_window:
            break
        is_speech = vad.is_speech(segment.tobytes(), sample_rate)
        if is_speech:
            segments.append(segment)

    if segments:
        segments_audio = np.concatenate(segments)
        output_segment = AudioSegment(
            data=segments_audio.tobytes(),
            sample_width=audio.sample_width,
            frame_rate=sample_rate,
            channels=num_channels
        )
        filtered_audio_path = audio_chunk_path.replace(".mp3", "_filtered.wav")
        output_segment.export(filtered_audio_path, format="wav")
        return filtered_audio_path
    else:
        return None

def transcribe_audio(audio_chunk):
    global total_api_time
    for attempt in range(3):
        try:
            global timer_running
            timer_running = True
            timer_thread = threading.Thread(target=live_timer)
            timer_thread.start()

            filtered_audio_path = filter_audio_segments(audio_chunk)
            if not filtered_audio_path:
                return ""

            with wave.open(filtered_audio_path, "rb") as audio_file:
                transcript = whisper_model.transcribe(audio_file)

            timer_running = False
            timer_thread.join()

            processing_time = time.time() - start_time
            total_api_time += processing_time  # Accumulate the API call time

            return transcript["text"]
        except Exception as e:
            print(f"Error transcribing {audio_chunk}: {e}. Retrying ({attempt+1}/3)...")
            if attempt == 2:
                raise e

def summarize_text(text):
    text_chunks = split_text_into_chunks(text, max_tokens=3000)
    chunk_summaries = []
    global total_api_time

    for chunk in text_chunks:
        for attempt in range(3):
            try:
                global timer_running
                timer_running = True
                timer_thread = threading.Thread(target=live_timer)
                timer_thread.start()

                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": f"Please summarize the following text:\n\n{chunk}"}]
                )

                timer_running = False
                timer_thread.join()

                processing_time = time.time() - start_time
                total_api_time += processing_time  # Accumulate the API call time
                chunk_summary = response['choices'][0]['message']['content']
                chunk_summaries.append(chunk_summary)
                break
            except openai.error.APIError as e:
                print(f"Error summarizing text chunk: {e}. Retrying ({attempt+1}/3)...")
                if attempt == 2:
                    raise e

    final_summary = " ".join(chunk_summaries)
    return final_summary

def integrate_visual_and_text(transcript, frame_descriptions):
    integrated_content = []
    transcript_lines = transcript.splitlines()

    for i, frame_description in enumerate(frame_descriptions):
        corresponding_text = transcript_lines[i] if i < len(transcript_lines) else ""
        integrated_content.append(f"Frame Description: {frame_description}\nCorresponding Text: {corresponding_text}")

    return "\n".join(integrated_content)

def generate_title_from_summary(summary):
    prompt = f"Based on the following summary, please generate a concise title (no more than 5 words):\n\n{summary}"
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    title = sanitize_title(response['choices'][0]['message']['content'].strip())
    return title

def create_sop_content(procedures):
    sop_content = {
        "Document Control": {
            "Version": "1.0",
            "Approval Date": datetime.datetime.now().strftime("%m/%d/%Y"),
            "Approved by": "Nathaniel Ruiz – IT Manager & WMS Administrator",
            "Next Review Date": (datetime.datetime.now() + datetime.timedelta(days=365)).strftime("%m/%d/%Y")
        },
        "Purpose": "This SOP outlines the procedures based on the analysis of the summarized content.",
        "Scope": "This procedure applies to all relevant personnel as indicated by the video content.",
        "Responsibilities": [
            "Individuals involved in the process outlined by the video content.",
            "Supervisors or leads responsible for overseeing the process."
        ],
        "Procedure": procedures,
        "Review and Updates": "This SOP should be reviewed annually or after any significant changes. Updates must be documented and approved.",
        "Compliance": "Adherence to this SOP is mandatory. Non-compliance may result in corrective actions."
    }
    return sop_content

def write_sop_to_word(output_folder, sop_content):
    doc = Document()
    doc.add_heading('Standard Operating Procedure (SOP)', 0)

    doc.add_heading('Document Control', level=1)
    for key, value in sop_content["Document Control"].items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('Purpose', level=1)
    doc.add_paragraph(sop_content["Purpose"])

    doc.add_heading('Scope', level=1)
    doc.add_paragraph(sop_content["Scope"])

    doc.add_heading('Responsibilities', level=1)
    for responsibility in sop_content["Responsibilities"]:
        doc.add_paragraph(f"• {responsibility}")

    doc.add_heading('Procedure', level=1)
    for step in sop_content["Procedure"]:
        doc.add_paragraph(f"{step}")

    doc.add_heading('Review and Updates', level=1)
    doc.add_paragraph(sop_content["Review and Updates"])

    doc.add_heading('Compliance', level=1)
    doc.add_paragraph(sop_content["Compliance"])

    sop_filename = os.path.join(output_folder, 'SOP.docx')
    try:
        doc.save(sop_filename)
        print(f"SOP saved as {sop_filename}")
    except Exception as e:
        print(f"Failed to save SOP: {e}")

def process_existing_chunks(audio_chunks_folder):
    transcripts = []
    for audio_chunk in os.listdir(audio_chunks_folder):
        audio_chunk_path = os.path.join(audio_chunks_folder, audio_chunk)
        try:
            transcript = transcribe_audio(audio_chunk_path)
            transcripts.append(transcript)
        except Exception as e:
            print(f"Failed to transcribe {audio_chunk_path}: {e}. Skipping...")

    full_transcript = "\n".join(transcripts)
    return full_transcript

def main():
    if os.listdir(input_path):
        process_input = input("Would you like to process the files in the input folder? (Y/N): ").strip().lower()
        if process_input == "y":
            for video_file in os.listdir(input_path):
                if video_file.endswith(".mkv"):
                    original_file_path = os.path.join(input_path, video_file)

                    temp_folder_title = "Processing Summary"
                    output_folder, video_chunks_folder, audio_chunks_folder = create_output_folder(temp_folder_title)

                    video_chunks = split_video(original_file_path, video_chunks_folder, audio_chunks_folder)

                    full_transcript = process_existing_chunks(audio_chunks_folder)

                    shutil.move(original_file_path, os.path.join(archive_path, video_file))

                    create_sop = input("Would you like to create an SOP for this video? (Y/N): ").strip().lower()
                    if create_sop == "y":
                        summary = summarize_text(full_transcript)
                        frame_output_folder = os.path.join(output_folder, "Key Frames")
                        frame_paths = extract_key_frames(original_file_path, frame_output_folder)
                        frame_descriptions = extract_text_from_frames(frame_paths)
                        procedures = analyze_transcript(full_transcript, frame_descriptions)

                        sop_content = create_sop_content(procedures)
                        write_sop_to_word(output_folder, sop_content)

                        concise_title = generate_title_from_summary(summary)
                        new_folder_title = f"{concise_title} - {datetime.datetime.now().strftime('%Y-%m-%d')} - Project"
                        new_output_folder = os.path.join(output_path, new_folder_title)

                        try:
                            os.rename(output_folder, new_output_folder)
                        except OSError as e:
                            print(f"Error renaming folder: {e}")

            print(f"Total API time: {total_api_time:.2f} seconds")
        else:
            process_failed = input("Would you like to attempt to fix any failed transcriptions or summarizations? (Y/N): ").strip().lower()
            if process_failed == "y":
                output_folders = [f.path for f in os.scandir(output_path) if f.is_dir()]
                for folder in output_folders:
                    video_chunks_folder = os.path.join(folder, "Video Chunks")
                    audio_chunks_folder = os.path.join(folder, "Audio Chunks")
                    full_transcript = process_existing_chunks(audio_chunks_folder)
                    summary = summarize_text(full_transcript)
                    frame_output_folder = os.path.join(folder, "Key Frames")
                    frame_paths = extract_key_frames(folder, frame_output_folder)
                    frame_descriptions = extract_text_from_frames(frame_paths)
                    procedures = analyze_transcript(full_transcript, frame_descriptions)

                    sop_content = create_sop_content(procedures)
                    write_sop_to_word(folder, sop_content)

                    concise_title = generate_title_from_summary(summary)
                    new_folder_title = f"{concise_title} - {datetime.datetime.now().strftime('%Y-%m-%d')} - Project"
                    new_output_folder = os.path.join(output_path, new_folder_title)

                    try:
                        os.rename(folder, new_output_folder)
                    except OSError as e:
                        print(f"Error renaming folder: {e}")
            else:
                create_sop = input("Would you like to create an SOP for any of the files in the output folder? (Y/N): ").strip().lower()
                if create_sop == "y":
                    output_folders = [f for f in os.listdir(output_path) if os.path.isdir(os.path.join(output_path, f))]

                    if not output_folders:
                        print("No files available in the output folder.")
                        return

                    print("Available processed files:")
                    for idx, folder in enumerate(output_folders, 1):
                        print(f"{idx}. {folder}")

                    selected_option = input("Select the number of the file you want to create an SOP for (or 0 to cancel): ").strip()

                    if selected_option == "0":
                        print("Operation cancelled.")
                        return

                    try:
                        selected_idx = int(selected_option) - 1
                        if selected_idx < 0 or selected_idx >= len(output_folders):
                            print("Invalid selection. Operation cancelled.")
                            return
                    except ValueError:
                        print("Invalid input. Operation cancelled.")
                        return

                    selected_folder = output_folders[selected_idx]
                    selected_folder_path = os.path.join(output_path, selected_folder)

                    print("The SOP feature will analyze the text-based summary and transcript to create the SOP. Ensure the content is clear for AI to generate meaningful SOPs.")
                    proceed = input("Would you like to proceed? (Y/N): ").strip().lower()

                    if proceed != "y":
                        print("Operation cancelled.")
                        return

                    video_chunks_folder = os.path.join(selected_folder_path, "Video Chunks")
                    audio_chunks_folder = os.path.join(selected_folder_path, "Audio Chunks")

                    full_transcript = process_existing_chunks(audio_chunks_folder)
                    summary = summarize_text(full_transcript)
                    frame_output_folder = os.path.join(selected_folder_path, "Key Frames")
                    frame_paths = extract_key_frames(selected_folder_path, frame_output_folder)
                    frame_descriptions = extract_text_from_frames(frame_paths)
                    procedures = analyze_transcript(full_transcript, frame_descriptions)

                    sop_content = create_sop_content(procedures)
                    write_sop_to_word(selected_folder_path, sop_content)

                    concise_title = generate_title_from_summary(summary)
                    new_folder_title = f"{concise_title} - {datetime.datetime.now().strftime('%Y-%m-%d')} - Project"
                    new_output_folder = os.path.join(output_path, new_folder_title)

                    try:
                        os.rename(selected_folder_path, new_output_folder)
                    except OSError as e:
                        print(f"Error renaming folder: {e}")

            print(f"Total API time: {total_api_time:.2f} seconds")

if __name__ == "__main__":
    main()
load_model
from pydub import AudioSegment
import numpy as np
import webrtcvad
import wave
import contextlib

# Set your OpenAI API key
openai.api_key = "YOUR_API_KEY_HERE"

# Initialize CLIP model for visual analysis
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
whisper_model = load_model("base")  # Corrected: Load Whisper model for transcription

# Paths
input_path = r"C:\Users\it\Videos\Video Sum\Input"
archive_path = r"C:\Users\it\Videos\Video Sum\Input\Archive"
output_path = r"C:\Users\it\Videos\Video Sum\Output"

# Global variables for the timer
timer_running = False
start_time = 0
total_api_time = 0  # Added to accumulate the total API call time


def sanitize_title(title):
    sanitized_title = re.sub(r'[<>:"/\\|?*]', '', title).strip()
    return sanitized_title


def create_output_folder(title):
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    folder_name = f"{title} - {current_date} - Project"
    folder_path = os.path.join(output_path, folder_name)

    video_chunks_folder = os.path.join(folder_path, "Video Chunks")
    audio_chunks_folder = os.path.join(folder_path, "Audio Chunks")
    os.makedirs(video_chunks_folder, exist_ok=True)
    os.makedirs(audio_chunks_folder, exist_ok=True)

    return folder_path, video_chunks_folder, audio_chunks_folder


def split_video_chunk(chunk, video_output_folder, audio_output_folder, index):
    base_filename = os.path.splitext(os.path.basename(chunk.filename))[0]
    video_chunk_filename = os.path.join(video_output_folder, f"{base_filename}_chunk_{index + 1}.mp4")
    audio_chunk_filename = os.path.join(audio_output_folder, f"{base_filename}_chunk_{index + 1}.mp3")

    chunk = chunk.without_audio()
    chunk.write_videofile(video_chunk_filename, codec="libx264", bitrate="500k")
    chunk.audio.write_audiofile(audio_chunk_filename, codec="mp3", bitrate="128k")

    return video_chunk_filename, audio_chunk_filename


def split_video(input_file, video_output_folder, audio_output_folder, max_duration=600):
    clip = VideoFileClip(input_file)
    duration = clip.duration

    with ThreadPoolExecutor() as executor:
        futures = []
        for i in range(0, int(duration), max_duration):
            chunk = clip.subclip(i, min(i + max_duration, duration))
            futures.append(
                executor.submit(split_video_chunk, chunk, video_output_folder, audio_output_folder, i // max_duration))

        video_chunks = [future.result() for future in futures]

    return video_chunks


def extract_key_frames(video_path, output_folder, interval=60):
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS)
    frame_interval = int(fps * interval)

    os.makedirs(output_folder, exist_ok=True)
    frame_count = 0
    success, frame = cap.read()
    key_frame_paths = []

    while success:
        if frame_count % frame_interval == 0:
            frame_time = frame_count / fps
            frame_filename = f"frame_{int(frame_time)}.jpg"
            frame_path = os.path.join(output_folder, frame_filename)
            cv2.imwrite(frame_path, frame)
            key_frame_paths.append(frame_path)
        success, frame = cap.read()
        frame_count += 1

    cap.release()
    return key_frame_paths


def extract_text_from_frames(frame_paths):
    extracted_texts = []
    for frame_path in frame_paths:
        image = Image.open(frame_path)
        text = pytesseract.image_to_string(image)
        extracted_texts.append(text)
    return extracted_texts


def analyze_transcript(transcript, extracted_texts):
    procedures = []
    for line in transcript.splitlines():
        if "log in" in line.lower() or "login" in line.lower():
            procedures.append("Step: Log in to VSR system")
        elif "submit" in line.lower():
            procedures.append("Step: Submit the form or data")
        elif "process" in line.lower():
            procedures.append("Step: Process the order or request")
    for text in extracted_texts:
        if "order" in text.lower():
            procedures.append("Step: Review and finalize the order in the VSR system")
    if not procedures:
        procedures.append("No explicit procedures were found in the video.")
    return procedures


def live_timer():
    global timer_running, start_time, total_api_time
    start_time = time.time()
    while timer_running:
        elapsed_time = time.time() - start_time
        print(f"\rTime elapsed: {int(elapsed_time)} seconds", end="")
        time.sleep(1)


def split_text_into_chunks(text, max_tokens=3000):
    sentences = text.split('. ')
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_tokens:
            current_chunk += sentence + ". "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


def filter_audio_segments(audio_chunk_path):
    vad = webrtcvad.Vad()
    vad.set_mode(1)  # Set VAD aggressiveness (0-3). Higher means more aggressive filtering of noise.

    audio = AudioSegment.from_file(audio_chunk_path)
    raw_audio = np.array(audio.get_array_of_samples())

    sample_rate = audio.frame_rate
    num_channels = audio.channels

    window_duration = 0.03  # 30 ms windows
    samples_per_window = int(sample_rate * window_duration * num_channels)

    segments = []

    for i in range(0, len(raw_audio), samples_per_window):
        segment = raw_audio[i:i + samples_per_window]
        if len(segment) < samples_per_window:
            break
        is_speech = vad.is_speech(segment.tobytes(), sample_rate)
        if is_speech:
            segments.append(segment)

    if segments:
        segments_audio = np.concatenate(segments)
        output_segment = AudioSegment(
            data=segments_audio.tobytes(),
            sample_width=audio.sample_width,
            frame_rate=sample_rate,
            channels=num_channels
        )
        filtered_audio_path = audio_chunk_path.replace(".mp3", "_filtered.wav")
        output_segment.export(filtered_audio_path, format="wav")
        return filtered_audio_path
    else:
        return None


def transcribe_audio(audio_chunk):
    global total_api_time
    for attempt in range(3):
        try:
            global timer_running
            timer_running = True
            timer_thread = threading.Thread(target=live_timer)
            timer_thread.start()

            filtered_audio_path = filter_audio_segments(audio_chunk)
            if not filtered_audio_path:
                return ""

            with wave.open(filtered_audio_path, "rb") as audio_file:
                transcript = whisper_model.transcribe(audio_file)

            timer_running = False
            timer_thread.join()

            processing_time = time.time() - start_time
            total_api_time += processing_time  # Accumulate the API call time

            return transcript["text"]
        except Exception as e:
            print(f"Error transcribing {audio_chunk}: {e}. Retrying ({attempt + 1}/3)...")
            if attempt == 2:
                raise e


def summarize_text(text):
    text_chunks = split_text_into_chunks(text, max_tokens=3000)
    chunk_summaries = []
    global total_api_time

    for chunk in text_chunks:
        for attempt in range(3):
            try:
                global timer_running
                timer_running = True
                timer_thread = threading.Thread(target=live_timer)
                timer_thread.start()

                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": f"Please summarize the following text:\n\n{chunk}"}]
                )

                timer_running = False
                timer_thread.join()

                processing_time = time.time() - start_time
                total_api_time += processing_time  # Accumulate the API call time
                chunk_summary = response['choices'][0]['message']['content']
                chunk_summaries.append(chunk_summary)
                break
            except openai.error.APIError as e:
                print(f"Error summarizing text chunk: {e}. Retrying ({attempt + 1}/3)...")
                if attempt == 2:
                    raise e

    final_summary = " ".join(chunk_summaries)
    return final_summary


def integrate_visual_and_text(transcript, frame_descriptions):
    integrated_content = []
    transcript_lines = transcript.splitlines()

    for i, frame_description in enumerate(frame_descriptions):
        corresponding_text = transcript_lines[i] if i < len(transcript_lines) else ""
        integrated_content.append(f"Frame Description: {frame_description}\nCorresponding Text: {corresponding_text}")

    return "\n".join(integrated_content)


def generate_title_from_summary(summary):
    prompt = f"Based on the following summary, please generate a concise title (no more than 5 words):\n\n{summary}"
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    title = sanitize_title(response['choices'][0]['message']['content'].strip())
    return title


def create_sop_content(procedures):
    sop_content = {
        "Document Control": {
            "Version": "1.0",
            "Approval Date": datetime.datetime.now().strftime("%m/%d/%Y"),
            "Approved by": "Nathaniel Ruiz – IT Manager & WMS Administrator",
            "Next Review Date": (datetime.datetime.now() + datetime.timedelta(days=365)).strftime("%m/%d/%Y")
        },
        "Purpose": "This SOP outlines the procedures based on the analysis of the summarized content.",
        "Scope": "This procedure applies to all relevant personnel as indicated by the video content.",
        "Responsibilities": [
            "Individuals involved in the process outlined by the video content.",
            "Supervisors or leads responsible for overseeing the process."
        ],
        "Procedure": procedures,
        "Review and Updates": "This SOP should be reviewed annually or after any significant changes. Updates must be documented and approved.",
        "Compliance": "Adherence to this SOP is mandatory. Non-compliance may result in corrective actions."
    }
    return sop_content


def write_sop_to_word(output_folder, sop_content):
    doc = Document()
    doc.add_heading('Standard Operating Procedure (SOP)', 0)

    doc.add_heading('Document Control', level=1)
    for key, value in sop_content["Document Control"].items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading('Purpose', level=1)
    doc.add_paragraph(sop_content["Purpose"])

    doc.add_heading('Scope', level=1)
    doc.add_paragraph(sop_content["Scope"])

    doc.add_heading('Responsibilities', level=1)
    for responsibility in sop_content["Responsibilities"]:
        doc.add_paragraph(f"• {responsibility}")

    doc.add_heading('Procedure', level=1)
    for step in sop_content["Procedure"]:
        doc.add_paragraph(f"{step}")

    doc.add_heading('Review and Updates', level=1)
    doc.add_paragraph(sop_content["Review and Updates"])

    doc.add_heading('Compliance', level=1)
    doc.add_paragraph(sop_content["Compliance"])

    sop_filename = os.path.join(output_folder, 'SOP.docx')
    try:
        doc.save(sop_filename)
        print(f"SOP saved as {sop_filename}")
    except Exception as e:
        print(f"Failed to save SOP: {e}")


def process_existing_chunks(audio_chunks_folder):
    transcripts = []
    for audio_chunk in os.listdir(audio_chunks_folder):
        audio_chunk_path = os.path.join(audio_chunks_folder, audio_chunk)
        try:
            transcript = transcribe_audio(audio_chunk_path)
            transcripts.append(transcript)
        except Exception as e:
            print(f"Failed to transcribe {audio_chunk_path}: {e}. Skipping...")

    full_transcript = "\n".join(transcripts)
    return full_transcript


def main():
    if os.listdir(input_path):
        process_input = input("Would you like to process the files in the input folder? (Y/N): ").strip().lower()
        if process_input == "y":
            for video_file in os.listdir(input_path):
                if video_file.endswith(".mkv"):
                    original_file_path = os.path.join(input_path, video_file)

                    temp_folder_title = "Processing Summary"
                    output_folder, video_chunks_folder, audio_chunks_folder = create_output_folder(temp_folder_title)

                    video_chunks = split_video(original_file_path, video_chunks_folder, audio_chunks_folder)

                    full_transcript = process_existing_chunks(audio_chunks_folder)

                    shutil.move(original_file_path, os.path.join(archive_path, video_file))

                    create_sop = input("Would you like to create an SOP for this video? (Y/N): ").strip().lower()
                    if create_sop == "y":
                        summary = summarize_text(full_transcript)
                        frame_output_folder = os.path.join(output_folder, "Key Frames")
                        frame_paths = extract_key_frames(original_file_path, frame_output_folder)
                        frame_descriptions = extract_text_from_frames(frame_paths)
                        procedures = analyze_transcript(full_transcript, frame_descriptions)

                        sop_content = create_sop_content(procedures)
                        write_sop_to_word(output_folder, sop_content)

                        concise_title = generate_title_from_summary(summary)
                        new_folder_title = f"{concise_title} - {datetime.datetime.now().strftime('%Y-%m-%d')} - Project"
                        new_output_folder = os.path.join(output_path, new_folder_title)

                        try:
                            os.rename(output_folder, new_output_folder)
                        except OSError as e:
                            print(f"Error renaming folder: {e}")

            print(f"Total API time: {total_api_time:.2f} seconds")
        else:
            process_failed = input(
                "Would you like to attempt to fix any failed transcriptions or summarizations? (Y/N): ").strip().lower()
            if process_failed == "y":
                output_folders = [f.path for f in os.scandir(output_path) if f.is_dir()]
                for folder in output_folders:
                    video_chunks_folder = os.path.join(folder, "Video Chunks")
                    audio_chunks_folder = os.path.join(folder, "Audio Chunks")
                    full_transcript = process_existing_chunks(audio_chunks_folder)
                    summary = summarize_text(full_transcript)
                    frame_output_folder = os.path.join(folder, "Key Frames")
                    frame_paths = extract_key_frames(folder, frame_output_folder)
                    frame_descriptions = extract_text_from_frames(frame_paths)
                    procedures = analyze_transcript(full_transcript, frame_descriptions)

                    sop_content = create_sop_content(procedures)
                    write_sop_to_word(folder, sop_content)

                    concise_title = generate_title_from_summary(summary)
                    new_folder_title = f"{concise_title} - {datetime.datetime.now().strftime('%Y-%m-%d')} - Project"
                    new_output_folder = os.path.join(output_path, new_folder_title)

                    try:
                        os.rename(folder, new_output_folder)
                    except OSError as e:
                        print(f"Error renaming folder: {e}")
            else:
                create_sop = input(
                    "Would you like to create an SOP for any of the files in the output folder? (Y/N): ").strip().lower()
                if create_sop == "y":
                    output_folders = [f for f in os.listdir(output_path) if os.path.isdir(os.path.join(output_path, f))]

                    if not output_folders:
                        print("No files available in the output folder.")
                        return

                    print("Available processed files:")
                    for idx, folder in enumerate(output_folders, 1):
                        print(f"{idx}. {folder}")

                    selected_option = input(
                        "Select the number of the file you want to create an SOP for (or 0 to cancel): ").strip()

                    if selected_option == "0":
                        print("Operation cancelled.")
                        return

                    try:
                        selected_idx = int(selected_option) - 1
                        if selected_idx < 0 or selected_idx >= len(output_folders):
                            print("Invalid selection. Operation cancelled.")
                            return
                    except ValueError:
                        print("Invalid input. Operation cancelled.")
                        return

                    selected_folder = output_folders[selected_idx]
                    selected_folder_path = os.path.join(output_path, selected_folder)

                    print(
                        "The SOP feature will analyze the text-based summary and transcript to create the SOP. Ensure the content is clear for AI to generate meaningful SOPs.")
                    proceed = input("Would you like to proceed? (Y/N): ").strip().lower()

                    if proceed != "y":
                        print("Operation cancelled.")
                        return

                    video_chunks_folder = os.path.join(selected_folder_path, "Video Chunks")
                    audio_chunks_folder = os.path.join(selected_folder_path, "Audio Chunks")

                    full_transcript = process_existing_chunks(audio_chunks_folder)
                    summary = summarize_text(full_transcript)
                    frame_output_folder = os.path.join(selected_folder_path, "Key Frames")
                    frame_paths = extract_key_frames(selected_folder_path, frame_output_folder)
                    frame_descriptions = extract_text_from_frames(frame_paths)
                    procedures = analyze_transcript(full_transcript, frame_descriptions)

                    sop_content = create_sop_content(procedures)
                    write_sop_to_word(selected_folder_path, sop_content)

                    concise_title = generate_title_from_summary(summary)
                    new_folder_title = f"{concise_title} - {datetime.datetime.now().strftime('%Y-%m-%d')} - Project"
                    new_output_folder = os.path.join(output_path, new_folder_title)

                    try:
                        os.rename(selected_folder_path, new_output_folder)
                    except OSError as e:
                        print(f"Error renaming folder: {e}")

            print(f"Total API time: {total_api_time:.2f} seconds")


if __name__ == "__main__":
    main()
